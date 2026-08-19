"""Convert docs/SYSTEM-DOCUMENTATION.md → Word .docx"""
from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
SRC = ROOT / "SYSTEM-DOCUMENTATION.md"
OUT = ROOT / "IT-Dir-TechStores-System-Documentation.docx"


def shade_cell(cell, hex_color="1A365D"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_formatted_inline(paragraph, content):
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))", content)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        elif part.startswith("[") and "](" in part:
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            if m:
                run = paragraph.add_run(f"{m.group(1)} ({m.group(2)})")
                run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
            else:
                paragraph.add_run(part)
        else:
            paragraph.add_run(part)


def main():
    text = SRC.read_text(encoding="utf-8")
    text = re.sub(
        r"```mermaid[\s\S]*?```",
        "\n> *[Diagram omitted in Word — open docs/SYSTEM-DIAGRAMS.html for interactive Mermaid diagrams]*\n",
        text,
    )

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    for i in range(1, 4):
        try:
            h = styles[f"Heading {i}"]
            h.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
            h.font.name = "Calibri"
        except KeyError:
            pass

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RESTRICTED")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    h = doc.add_heading("IT Dir TechStores Information System", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "System Documentation — Architecture · Database · DFD · Flowcharts · Modules"
    )
    r.italic = True
    r.font.size = Pt(11)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        "Cost Centre Z04P2SP212  ·  From docs/SYSTEM-DOCUMENTATION.md  ·  28 July 2026"
    )
    mr.font.size = Pt(9)
    mr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

    lines = text.splitlines()
    i = 0
    in_code = False
    code_buf = []
    table_rows = []
    skipped_first_h1 = False

    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                block = "\n".join(code_buf)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run(block if block else " ")
                run.font.name = "Consolas"
                run.font.size = Pt(8.5)
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "F3F7FB")
                shd.set(qn("w:val"), "clear")
                p._p.get_or_add_pPr().append(shd)
                code_buf = []
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if "|" in line and line.strip().startswith("|"):
            if re.match(r"^\s*\|?\s*:?-{3,}", line.strip()):
                i += 1
                continue
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            table_rows.append(cells)
            nxt = lines[i + 1] if i + 1 < len(lines) else ""
            if not (nxt.strip().startswith("|") and "|" in nxt):
                if table_rows:
                    cols = max(len(r) for r in table_rows)
                    tbl = doc.add_table(rows=len(table_rows), cols=cols)
                    tbl.style = "Table Grid"
                    for r_idx, row in enumerate(table_rows):
                        for c_idx in range(cols):
                            cell = tbl.rows[r_idx].cells[c_idx]
                            val = row[c_idx] if c_idx < len(row) else ""
                            cell.text = ""
                            p = cell.paragraphs[0]
                            val_clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", val)
                            val_clean = re.sub(r"`([^`]+)`", r"\1", val_clean)
                            run = p.add_run(val_clean)
                            run.font.size = Pt(9)
                            run.font.name = "Calibri"
                            if r_idx == 0:
                                shade_cell(cell, "1A365D")
                                run.bold = True
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    doc.add_paragraph()
                table_rows = []
            i += 1
            continue

        hm = re.match(r"^(#{1,6})\s+(.*)$", line)
        if hm:
            level = len(hm.group(1))
            title_text = hm.group(2).strip()
            title_text = re.sub(r"\s*\{#.*\}\s*$", "", title_text)
            if level == 1 and not skipped_first_h1:
                skipped_first_h1 = True
                i += 1
                continue
            doc.add_heading(title_text, level=min(level, 3))
            i += 1
            continue

        if re.match(r"^---+$", line.strip()):
            i += 1
            continue

        if line.startswith(">"):
            content = re.sub(r"^>\s?", "", line)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            add_formatted_inline(p, content)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x44, 0x55, 0x66)
            i += 1
            continue

        if re.match(r"^[-*]\s+", line):
            content = re.sub(r"^[-*]\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_inline(p, content)
            i += 1
            continue

        if re.match(r"^\d+\.\s+", line):
            content = re.sub(r"^\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            add_formatted_inline(p, content)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        p = doc.add_paragraph()
        add_formatted_inline(p, line.strip())
        i += 1

    doc.add_paragraph()
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    er = end.add_run("RESTRICTED — End of System Documentation")
    er.bold = True
    er.font.size = Pt(10)
    er.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run(
        "Edit docs/SYSTEM-DOCUMENTATION.md then run: python docs/md_to_docx.py"
    )
    nr.font.size = Pt(8)
    nr.italic = True

    doc.save(OUT)
    print(f"Wrote {OUT}")
    print(f"Size: {OUT.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
