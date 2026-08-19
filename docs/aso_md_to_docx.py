"""Convert docs/ASO-COMPLIANCE-MATRIX.md → Word .docx"""
from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
SRC = ROOT / "ASO-COMPLIANCE-MATRIX.md"
OUT = ROOT / "IT-Dir-TechStores-ASO-Compliance-Matrix.docx"


def shade_cell(cell, hex_color="1A365D"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_inline(paragraph, content):
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", content)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def main():
    text = SRC.read_text(encoding="utf-8")
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)
    for i in range(1, 4):
        try:
            styles[f"Heading {i}"].font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        except KeyError:
            pass

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RESTRICTED")
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    h = doc.add_heading("TechStores ↔ Accounting Standing Orders", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Compliance Matrix — IT Dir Tech Stores · ASO August 2011")
    sr.italic = True

    lines = text.splitlines()
    i = 0
    skipped_h1 = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                i += 1
            i += 1
            continue

        if "|" in line and line.strip().startswith("|"):
            if re.match(r"^\s*\|?\s*:?-{3,}", line.strip()):
                i += 1
                continue
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            table_rows.append(cells)
            nxt = lines[i + 1] if i + 1 < len(lines) else ""
            if not (nxt.strip().startswith("|") and "|" in nxt):
                cols = max(len(r) for r in table_rows)
                tbl = doc.add_table(rows=len(table_rows), cols=cols)
                tbl.style = "Table Grid"
                for r_idx, row in enumerate(table_rows):
                    for c_idx in range(cols):
                        cell = tbl.rows[r_idx].cells[c_idx]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        val = row[c_idx] if c_idx < len(row) else ""
                        val = re.sub(r"\*\*([^*]+)\*\*", r"\1", val)
                        val = re.sub(r"`([^`]+)`", r"\1", val)
                        run = p.add_run(val)
                        run.font.size = Pt(8)
                        run.font.name = "Calibri"
                        if r_idx == 0:
                            shade_cell(cell)
                            run.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                doc.add_paragraph()
                table_rows = []
            i += 1
            continue

        hm = re.match(r"^(#{1,6})\s+(.*)$", line)
        if hm:
            level = len(hm.group(1))
            title_text = re.sub(r"\s*\{#.*\}\s*$", "", hm.group(2).strip())
            if level == 1 and not skipped_h1:
                skipped_h1 = True
                i += 1
                continue
            doc.add_heading(title_text, level=min(level, 3))
            i += 1
            continue

        if re.match(r"^---+$", line.strip()):
            i += 1
            continue

        if line.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            add_inline(p, re.sub(r"^>\s?", "", line))
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        if re.match(r"^[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, re.sub(r"^[-*]\s+", "", line))
            i += 1
            continue

        if re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s+", "", line))
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline(p, line.strip())
        i += 1

    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    er = end.add_run("RESTRICTED — End of ASO Compliance Matrix")
    er.bold = True
    er.font.size = Pt(9)

    doc.save(OUT)
    print(f"Wrote {OUT} ({OUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
